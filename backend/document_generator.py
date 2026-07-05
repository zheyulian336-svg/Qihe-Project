"""把合同文字转换成可下载的 .docx 文件。"""

import io

from docx import Document


def generate_contract_docx(contract_text: str, disclaimer: str = "") -> io.BytesIO:
    document = Document()
    title_set = False

    for raw_line in contract_text.split("\n"):
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if not title_set:
            document.add_heading(line, level=1)
            title_set = True
            continue
        document.add_paragraph(line)

    # 免责声明写在文档末尾
    if disclaimer:
        document.add_paragraph("")
        document.add_paragraph(disclaimer)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
